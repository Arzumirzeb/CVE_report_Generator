from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document
from markdown2 import markdown

def create_pdf(cve_info, file_path):
    """Generate a PDF report without tables and with a simple list for references."""
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    elements = []

    # Define styles
    styles = getSampleStyleSheet()
    heading_style = styles['Heading1']
    body_style = styles['BodyText']

    # Title
    title = f"CVE Report for {cve_info['cve_title']}"
    elements.append(Paragraph(title, heading_style))
    elements.append(Spacer(1, 0.2*inch))  # Add space after the title

    # CVSS Score and Vector
    cvss_score_vector = f"CVSS Score: {cve_info['cvss_score']} | CVSS Vector: {cve_info['cvss_vector']}"
    elements.append(Paragraph(cvss_score_vector, body_style))
    
    # State
    state = f"State: {cve_info.get('state', 'N/A')}"
    elements.append(Paragraph(state, body_style))
    elements.append(Spacer(1, 0.4*inch))  # Add more space before the description

    # Description
    description = f"Description: {cve_info['description']}"
    elements.append(Paragraph(description, body_style))
    elements.append(Spacer(1, 0.4*inch))  # Add space before the affected assets

    # Affected Assets
    if cve_info.get('affected_assets'):
        assets_header = Paragraph("Affected Assets:", heading_style)
        elements.append(assets_header)
        for asset in cve_info['affected_assets']:
            asset_text = f"Vendor: {asset.get('vendor', 'N/A')}, Product: {asset.get('product', 'N/A')}"
            elements.append(Paragraph(asset_text, body_style))
        elements.append(Spacer(1, 0.4*inch))  # Add space before the exploits

    # Exploits
    if cve_info.get('exploits'):
        exploits_header = Paragraph("Exploits:", heading_style)
        elements.append(exploits_header)
        for exploit in cve_info['exploits']:
            exploit_text = f"Title: {exploit.get('title', 'N/A')}, Verified: {exploit.get('verified', 'N/A')}"
            elements.append(Paragraph(exploit_text, body_style))
        elements.append(Spacer(1, 0.4*inch))  # Add space before references

    # References
    if cve_info.get('references'):
        references_header = Paragraph("References:", heading_style)
        elements.append(references_header)
        for index, ref in enumerate(cve_info['references'], start=1):
            ref_text = f"{index}. {ref.get('link', 'N/A')}"
            elements.append(Paragraph(ref_text, body_style))

    # Build PDF
    doc.build(elements)


def create_docx(cve_info, file_path):
    """Generate a DOCX report."""
    doc = Document()

    # Title
    doc.add_heading(f"CVE Report for {cve_info['cve_title']}", level=1)

    # CVSS Score and Vector
    doc.add_paragraph(f"CVSS Score: {cve_info['cvss_score']} | CVSS Vector: {cve_info['cvss_vector']}")
    
    # State
    doc.add_paragraph(f"State: {cve_info.get('state', 'N/A')}")

    # Description
    doc.add_paragraph(f"Description: {cve_info['description']}")

    # Affected Assets
    if cve_info.get('affected_assets'):
        doc.add_heading('Affected Assets:', level=2)
        for asset in cve_info['affected_assets']:
            doc.add_paragraph(f"Vendor: {asset.get('vendor', 'N/A')}, Product: {asset.get('product', 'N/A')}")

    # Exploits (Just title and whether verified or not)
    if cve_info.get('exploits'):
        doc.add_heading('Exploits:', level=2)
        for exploit in cve_info['exploits']:
            doc.add_paragraph(f"Title: {exploit.get('title', 'N/A')}, Verified: {exploit.get('verified', 'N/A')}")

    # References (Just links, ordered)
    if cve_info.get('references'):
        doc.add_heading('References:', level=2)
        for index, ref in enumerate(cve_info['references'], start=1):
            doc.add_paragraph(f"{index}. {ref.get('link', 'N/A')}")

    # Save the DOCX
    doc.save(file_path)


def create_html(cve_info, file_path):
    """Generate an HTML report."""
    html_content = (
        "<html>"
        "<head><title>CVE Report for {0}</title></head>"
        "<body>"
        "<h1>CVE Report for {0}</h1>"
        "<p><strong>CVSS Score:</strong> {1} | <strong>CVSS Vector:</strong> {2}</p>"
        "<p><strong>Description:</strong> {3}</p>"
        "<h2>Affected Assets:</h2>"
        "<ul>"
        "{4}"
        "</ul>"
        "<h2>Exploits:</h2>"
        "<ul>"
        "{5}"
        "</ul>"
        "<h2>References:</h2>"
        "<ul>"
        "{6}"
        "</ul>"
        "</body>"
        "</html>"
    ).format(
        cve_info['cve_title'],
        cve_info['cvss_score'],
        cve_info['cvss_vector'],
        cve_info['description'],
        "".join(f"<li>Vendor: {asset['vendor']}, Product: {asset['product']}</li>" for asset in cve_info['affected_assets']),
        "".join(f"<li>Title: {exploit['title']} - Verified: {exploit['verified']} - Download Link: <a href='{exploit['download_link']}'>Download</a> - Exploit Link: <a href='{exploit['exploit_link']}'>Exploit</a></li>" for exploit in cve_info['exploits']),
        "".join(f"<li>Description: {ref['description']} - Link: <a href='{ref['link']}'>Reference</a></li>" for ref in cve_info['references'])
    )
    with open(file_path, "w") as file:
        file.write(html_content)

def create_md(cve_info, file_path):
    """Generate a Markdown report."""
    markdown_content = (
        "# CVE Report for {0}\n\n"
        "**CVSS Score:** {1} | **CVSS Vector:** {2}\n\n"
        "**Description:** {3}\n\n"
        "## Affected Assets:\n"
        "{4}\n\n"
        "## Exploits:\n"
        "{5}\n\n"
        "## References:\n"
        "{6}\n"
    ).format(
        cve_info['cve_title'],
        cve_info['cvss_score'],
        cve_info['cvss_vector'],
        cve_info['description'],
        "".join(f"- Vendor: {asset['vendor']}, Product: {asset['product']}\n" for asset in cve_info['affected_assets']),
        "".join(f"- Title: {exploit['title']} - Verified: {exploit['verified']} - Download Link: [{exploit['download_link']}]({exploit['download_link']}) - Exploit Link: [{exploit['exploit_link']}]({exploit['exploit_link']})\n" for exploit in cve_info['exploits']),
        "".join(f"- Description: {ref['description']} - Link: [{ref['link']}]({ref['link']})\n" for ref in cve_info['references'])
    )
    with open(file_path, "w") as file:
        file.write(markdown_content)

